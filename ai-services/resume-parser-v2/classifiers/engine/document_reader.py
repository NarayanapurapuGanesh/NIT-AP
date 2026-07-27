"""
Layer 2: Document Reader.
Format-specific reader for PDF, DOCX, DOC, and Image formats.
Extracts raw page text, document structure elements, and handles encrypted/corrupted files gracefully.
"""

import io
from typing import Any, Dict, List
from core.exceptions import BaseAppException
from core.logging import get_logger

logger = get_logger("document_reader")


class RawDocumentContent:
    def __init__(
        self,
        format_type: str,
        pages: List[str],
        metadata_raw: Dict[str, Any],
        is_scanned: bool = False,
        error_msg: str | None = None,
    ) -> None:
        self.format_type = format_type
        self.pages = pages
        self.metadata_raw = metadata_raw
        self.is_scanned = is_scanned
        self.error_msg = error_msg

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.pages)


class DocumentReader:
    """Layer 2: Format-specific reader engine."""

    def read_document(self, format_type: str, content: bytes) -> RawDocumentContent:
        """Reads document bytes safely and returns raw content data object."""
        try:
            if format_type == "pdf":
                return self._read_pdf(content)
            elif format_type == "docx":
                return self._read_docx(content)
            elif format_type == "doc":
                return self._read_doc(content)
            elif format_type in ["png", "jpeg", "tiff"]:
                return self._read_image(format_type, content)
            else:
                return self._read_fallback_text(format_type, content)
        except Exception as exc:
            logger.warning("Error reading document stream", format_type=format_type, error=str(exc))
            return RawDocumentContent(
                format_type=format_type,
                pages=[],
                metadata_raw={},
                error_msg=f"Failed to read file stream: {str(exc)}",
            )

    def _read_pdf(self, content: bytes) -> RawDocumentContent:
        import pypdf

        pages_text: List[str] = []
        raw_meta: Dict[str, Any] = {}
        is_scanned = False

        try:
            reader = pypdf.PdfReader(io.BytesIO(content))

            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    return RawDocumentContent(
                        format_type="pdf",
                        pages=[],
                        metadata_raw={},
                        error_msg="PDF is password encrypted and cannot be read.",
                    )

            if reader.metadata:
                for k, v in reader.metadata.items():
                    raw_meta[str(k).replace("/", "")] = str(v)

            total_chars = 0
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages_text.append(text)
                total_chars += len(text.strip())

            raw_meta["page_count"] = len(reader.pages)
            if total_chars < 50 and len(reader.pages) > 0:
                is_scanned = True

            return RawDocumentContent(
                format_type="pdf",
                pages=pages_text,
                metadata_raw=raw_meta,
                is_scanned=is_scanned,
            )

        except Exception as exc:
            try:
                text = content.decode("utf-8", errors="ignore")
                if text and any(c.isalnum() for c in text):
                    return RawDocumentContent(
                        format_type="pdf",
                        pages=[text],
                        metadata_raw={"page_count": 1, "fallback": True},
                    )
            except Exception:
                pass
            return RawDocumentContent(
                format_type="pdf",
                pages=[],
                metadata_raw={},
                error_msg=f"Corrupted or invalid PDF format: {str(exc)}",
            )

    def _read_docx(self, content: bytes) -> RawDocumentContent:
        import docx

        try:
            doc = docx.Document(io.BytesIO(content))
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            raw_meta = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            return RawDocumentContent(
                format_type="docx",
                pages=[full_text],
                metadata_raw=raw_meta,
            )
        except Exception as exc:
            return RawDocumentContent(
                format_type="docx",
                pages=[],
                metadata_raw={},
                error_msg=f"Invalid or corrupted DOCX file: {str(exc)}",
            )

    def _read_doc(self, content: bytes) -> RawDocumentContent:
        # Basic plain-text ASCII/UTF-8 extraction fallback for legacy .doc binary files
        try:
            text = content.decode("utf-8", errors="ignore")
            clean_text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
            return RawDocumentContent(
                format_type="doc",
                pages=[clean_text],
                metadata_raw={"binary_length": len(content)},
            )
        except Exception as exc:
            return RawDocumentContent(
                format_type="doc",
                pages=[],
                metadata_raw={},
                error_msg=f"Failed to extract binary DOC file: {str(exc)}",
            )

    def _read_image(self, format_type: str, content: bytes) -> RawDocumentContent:
        from PIL import Image

        try:
            img = Image.open(io.BytesIO(content))
            width, height = img.size
            raw_meta = {
                "image_width": width,
                "image_height": height,
                "image_format": img.format,
                "image_mode": img.mode,
            }
            return RawDocumentContent(
                format_type=format_type,
                pages=[],
                metadata_raw=raw_meta,
                is_scanned=True,
            )
        except Exception as exc:
            return RawDocumentContent(
                format_type=format_type,
                pages=[],
                metadata_raw={},
                error_msg=f"Invalid or corrupted image format: {str(exc)}",
            )

    def _read_fallback_text(self, format_type: str, content: bytes) -> RawDocumentContent:
        try:
            text = content.decode("utf-8", errors="ignore")
            return RawDocumentContent(
                format_type=format_type,
                pages=[text],
                metadata_raw={},
            )
        except Exception:
            return RawDocumentContent(
                format_type=format_type,
                pages=[],
                metadata_raw={},
                error_msg="Unreadable plain text format.",
            )
