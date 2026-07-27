"""
Component 3: Document Loaders & Plugin Factory.
Provides abstract IDocumentLoader interface and format-specific loader implementations.
"""

from abc import ABC, abstractmethod
import io
from typing import Any, Dict, List
import docx
import pypdf
from app.document.exceptions.ingestion_exceptions import ExtractionException
from core.logging import get_logger

logger = get_logger("document_loader")


class LoadedRawDocument:
    def __init__(
        self,
        format_type: str,
        pages_text: List[str],
        metadata_dict: Dict[str, Any],
        raw_stream: bytes,
    ) -> None:
        self.format_type = format_type
        self.pages_text = pages_text
        self.metadata_dict = metadata_dict
        self.raw_stream = raw_stream


class IDocumentLoader(ABC):
    """Abstract interface for format-specific document loaders."""

    @abstractmethod
    def load(self, content: bytes) -> LoadedRawDocument:
        pass


class PdfLoader(IDocumentLoader):
    """PDF document loader using pypdf and stream handlers."""

    def load(self, content: bytes) -> LoadedRawDocument:
        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages_text: List[str] = []
            meta: Dict[str, Any] = {}

            if reader.metadata:
                for k, v in reader.metadata.items():
                    meta[str(k).replace("/", "")] = str(v)

            meta["page_count"] = len(reader.pages)

            for page in reader.pages:
                txt = page.extract_text() or ""
                pages_text.append(txt)

            return LoadedRawDocument(
                format_type="pdf",
                pages_text=pages_text,
                metadata_dict=meta,
                raw_stream=content,
            )
        except Exception as exc:
            try:
                text = content.decode("utf-8", errors="ignore")
                if text and any(c.isalnum() for c in text):
                    return LoadedRawDocument(
                        format_type="pdf",
                        pages_text=[text],
                        metadata_dict={"page_count": 1, "fallback": True},
                        raw_stream=content,
                    )
            except Exception:
                pass
            raise ExtractionException(f"PDF loading error: {str(exc)}") from exc


class DocxLoader(IDocumentLoader):
    """DOCX document loader using python-docx."""

    def load(self, content: bytes) -> LoadedRawDocument:
        try:
            doc = docx.Document(io.BytesIO(content))
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            meta = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "page_count": max(1, len(doc.paragraphs) // 30),
            }
            return LoadedRawDocument(
                format_type="docx",
                pages_text=[full_text],
                metadata_dict=meta,
                raw_stream=content,
            )
        except Exception as exc:
            raise ExtractionException(f"DOCX loading error: {str(exc)}") from exc


class TxtLoader(IDocumentLoader):
    """Plain text loader."""

    def load(self, content: bytes) -> LoadedRawDocument:
        try:
            text = content.decode("utf-8", errors="ignore")
            return LoadedRawDocument(
                format_type="txt",
                pages_text=[text],
                metadata_dict={"page_count": 1, "char_count": len(text)},
                raw_stream=content,
            )
        except Exception as exc:
            raise ExtractionException(f"Text loading error: {str(exc)}") from exc


class RtfLoader(IDocumentLoader):
    """RTF text loader."""

    def load(self, content: bytes) -> LoadedRawDocument:
        try:
            text = content.decode("utf-8", errors="ignore")
            clean_text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
            return LoadedRawDocument(
                format_type="rtf",
                pages_text=[clean_text],
                metadata_dict={"page_count": 1, "char_count": len(clean_text)},
                raw_stream=content,
            )
        except Exception as exc:
            raise ExtractionException(f"RTF loading error: {str(exc)}") from exc


class DocumentLoaderFactory:
    """Factory pattern for instantiating format loaders."""

    @staticmethod
    def get_loader(format_type: str) -> IDocumentLoader:
        if format_type == "pdf":
            return PdfLoader()
        elif format_type == "docx":
            return DocxLoader()
        elif format_type == "rtf":
            return RtfLoader()
        else:
            return TxtLoader()
